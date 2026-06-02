from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io

CASE = {
    "number": "9664",
    "year": "2022",
    "petitioner_zh": "方銘",
    "petitioner_en": "FONG, MING",
    "respondent_zh": "李慧儀",
    "respondent_en": "LEE, WAI YEE",
    "petitioner_address_zh": "屯門青盈路28號The Bloomsway第3座G座",
    "petitioner_address_en": "GA, Tower 3, The Bloomsway, 28 Tsing Ying Road, Tuen Mun, New Territories",
    "respondent_hkid": "Z278264(7)",
    "custody_order_date_zh": "2021年12月14日",
    "custody_order_date_en": "12 December 2023",
    "child1_zh": "方顯霖", "child1_en": "FONG, HIUN LING",
    "child2_zh": "方晴兒", "child2_en": "FONG, CHING YI",
}


def _new_doc():
    doc = Document()
    s = doc.sections[0]
    s.page_height = Cm(29.7)
    s.page_width = Cm(21.0)
    s.left_margin = Cm(3.2)
    s.right_margin = Cm(2.5)
    s.top_margin = Cm(2.5)
    s.bottom_margin = Cm(2.5)
    for p in doc.paragraphs:
        p.clear()
    return doc


def _zh_court_header(doc, title, case):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(
        f"香港特別行政區\n區域法院\n"
        f"婚姻訴訟案件編號  {case['year']}  年  {case['number']}  號"
    ).font.size = Pt(12)
    doc.add_paragraph()
    _party_lines(doc, case)
    _divider(doc)
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t.add_run(title).bold = True
    doc.add_paragraph()


def _party_lines(doc, case):
    p1 = doc.add_paragraph()
    p1.add_run(f"\t{case['petitioner_zh']}（{case['petitioner_en']}）\t\t\t呈請人")
    doc.add_paragraph()
    c = doc.add_paragraph()
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c.add_run("與")
    doc.add_paragraph()
    p2 = doc.add_paragraph()
    p2.add_run(f"\t{case['respondent_zh']}（{case['respondent_en']}）\t\t\t答辯人")


def _divider(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("_" * 50)


def _numbered(doc, n, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.left_indent = Cm(0.74)
    p.add_run(f"{n}.\t{text}")


def _body(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.add_run(text)
    return p


def _affirmation_footer_zh(doc, sign_year, sign_month, sign_day):
    doc.add_paragraph()
    doc.add_paragraph()
    sig = doc.add_paragraph()
    sig.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    sig.add_run("___________________________")
    role = doc.add_paragraph()
    role.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    role.add_run("（確認者簽署）")
    doc.add_paragraph()
    doc.add_paragraph(
        f"此項確認是於  {sign_year}  年  {sign_month}  月  {sign_day}  日\n"
        "在香港特別行政區\n區域法院作出。"
    )
    doc.add_paragraph()
    doc.add_paragraph("在本人面前作出，")
    doc.add_paragraph()
    doc.add_paragraph("司法機構監誓員：___________")


def _to_bytes(doc):
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def children_zh(data):
    case = data.get("case", CASE)
    if data.get("both_children", True):
        return f"{case['child1_zh']} 及 {case['child2_zh']}"
    return case["child1_zh"] if data.get("child") == "child1" else case["child2_zh"]


def children_en(data):
    case = data.get("case", CASE)
    if data.get("both_children", True):
        return f"{case['child1_en']} and {case['child2_en']}"
    return case["child1_en"] if data.get("child") == "child1" else case["child2_en"]


# ── Option A ─────────────────────────────────────────────────────────────────

def generate_zh_affirmation_a(data):
    """誓詞 — Option A (mother agrees)."""
    doc = _new_doc()
    case = data.get("case", CASE)
    kids = children_zh(data)
    _zh_court_header(doc, "誓詞", case)
    _body(doc, f"本人  {case['petitioner_zh']}  ，現居於  {case['petitioner_address_zh']}。")
    doc.add_paragraph()
    _body(doc, "謹此謹以至誠確認以下內容：")
    doc.add_paragraph()
    _numbered(doc, 1,
        f"我是本案的呈請人，依據  {case['custody_order_date_zh']}  的法庭命令，"
        f"呈請人及答辯人獲授予家庭中的子女  {kids}  的共同管養權。")
    _numbered(doc, 2,
        f"我申請帶該等子女  {kids}  暫時離開香港，"
        f"日期由  {data['start_date_zh']}  至  {data['end_date_zh']}，"
        f"目的是  {data['destination_zh']}旅遊。")
    _numbered(doc, 3,
        f"我現附上我的承諾書，日期是  {data['undertaking_date_zh']}，"
        f"及答辯人的同意書，日期是  {data['consent_date_zh']}，"
        f"並分別標記為證物 [A] 和 [B]。")
    _numbered(doc, 4,
        "我認明同意書內答辯人的簽名，我亦確信她完全明白該同意書的效力和後果，"
        "而該份同意書是出於自願的。")
    doc.add_paragraph()
    _body(doc, "本人謹此謹以至誠確認，此誓詞內容一概是真。")
    _affirmation_footer_zh(doc, data["sign_year"], data["sign_month"], data["sign_day"])
    return _to_bytes(doc)


def generate_zh_undertaking(data):
    """承諾書 — Options A and B."""
    doc = _new_doc()
    case = data.get("case", CASE)
    kids = children_zh(data)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(
        f"香港特別行政區\n區域法院\n"
        f"婚姻訴訟  {case['year']}  年第  {case['number']}  號"
    ).font.size = Pt(12)
    doc.add_paragraph()
    _party_lines(doc, case)
    _divider(doc)
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t.add_run("承諾書").bold = True
    doc.add_paragraph()
    _body(doc, "本人為這宗案件的呈請人。")
    doc.add_paragraph()
    _body(doc,
        f"現就本人意欲為了  {data['destination_zh']}旅遊  而攜帶家庭子女  {kids}  "
        f"由  {data['start_date_zh']}  至  {data['end_date_zh']}  或/及  "
        f"於上述子女年滿18歲前不時離開香港一事，本人現向法庭承諾，"
        f"若法庭要求本人把上述子女帶回香港，本人定當遵循。")
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph(f"此份承諾書於  {data['sign_date_zh']}  簽署")
    doc.add_paragraph()
    doc.add_paragraph()
    sig = doc.add_paragraph()
    sig.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    sig.add_run("___________________________")
    r = doc.add_paragraph()
    r.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r.add_run("呈請人（簽署）")
    return _to_bytes(doc)


def generate_zh_consent(data):
    """同意書 — Option A only (mother/respondent signs)."""
    doc = _new_doc()
    case = data.get("case", CASE)
    kids = children_zh(data)
    _zh_court_header(doc, "同意書", case)
    _body(doc, "本人為這宗案件的答辯人。")
    doc.add_paragraph()
    _body(doc,
        f"現就呈請人意欲為了  {data['destination_zh']}旅遊  而攜帶家庭子女  {kids}  "
        f"由  {data['start_date_zh']}  至  {data['end_date_zh']}  或/及  "
        f"於上述子女年滿18歲前不時離開香港一事，本人在此作出書面同意，"
        f"該等子女的父可攜帶上述子女在上述期間為上述理由離開香港。")
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph(f"此份同意書於  {case['year']}  年  ______  月  ______  日簽署")
    doc.add_paragraph()
    doc.add_paragraph()
    sig = doc.add_paragraph()
    sig.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    sig.add_run("___________________________")
    r = doc.add_paragraph()
    r.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r.add_run("答辯人（簽署）")
    return _to_bytes(doc)


def generate_zh_annex(data):
    """附錄一 cover page."""
    doc = _new_doc()
    case = data.get("case", CASE)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(
        f"香港特別行政區\n區域法院\n"
        f"婚姻訴訟案件編號  {case['year']}  年  {case['number']}  號"
    ).font.size = Pt(12)
    doc.add_paragraph()
    p1 = doc.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1.add_run(case["petitioner_zh"])
    doc.add_paragraph()
    c = doc.add_paragraph()
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c.add_run("與")
    doc.add_paragraph()
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.add_run(case["respondent_zh"])
    _divider(doc)
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t.add_run("附錄  一").bold = True
    return _to_bytes(doc)


# ── Option B ─────────────────────────────────────────────────────────────────

def generate_zh_affirmation_b(data):
    """誓詞 — Option B (mother did NOT sign consent)."""
    doc = _new_doc()
    case = data.get("case", CASE)
    kids = children_zh(data)
    _zh_court_header(doc, "誓詞", case)
    _body(doc, f"本人  {case['petitioner_zh']}  ，現居於  {case['petitioner_address_zh']}。")
    doc.add_paragraph()
    _body(doc, "謹此謹以至誠確認以下內容：")
    doc.add_paragraph()
    _numbered(doc, 1,
        f"我是本案的呈請人，依據  {case['custody_order_date_zh']}  的法庭命令，"
        f"呈請人及答辯人獲授予家庭中的子女  {kids}  的共同管養權。")
    _numbered(doc, 2,
        f"我申請獲得許可，可以帶該等子女  {kids}  "
        f"由  {data['start_date_zh']}  至  {data['end_date_zh']}  "
        f"暫時離開香港，脫離法庭司法管轄區，目的是  {data['destination_zh']}旅遊；")
    _numbered(doc, 3,
        f"我現附上我的承諾書，日期是  {data['undertaking_date_zh']}；")
    _numbered(doc, 4,
        f"答辯人於  {data['respondent_email_date_zh']}  {data['respondent_response_desc']}，"
        f"但未有簽署及交回法庭標準同意書。由於未取得正式簽署之同意書，"
        f"本人需向法庭申請命令；")
    _numbered(doc, 5, "相關細節，請見附錄一；")
    _numbered(doc, 6, "懇請法庭按照我申請的條款授予命令。")
    doc.add_paragraph()
    _body(doc, "本人謹此謹以至誠確認，此誓詞內容一概是真。")
    _affirmation_footer_zh(doc, data["sign_year"], data["sign_month"], data["sign_day"])
    return _to_bytes(doc)


def generate_zh_notice_b(data):
    """申請帶多名子女離境許可通知書 — Option B."""
    doc = _new_doc()
    case = data.get("case", CASE)
    kids = children_zh(data)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(
        f"香港特別行政區\n區域法院\n"
        f"婚姻訴訟  {case['year']}  年  {case['number']}  號"
    ).font.size = Pt(12)
    doc.add_paragraph()
    _party_lines(doc, case)
    _divider(doc)
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t.add_run("申請帶多名子女離境許可通知書").bold = True
    doc.add_paragraph()
    addr = doc.add_paragraph()
    addr.add_run(f"致：  {case['respondent_zh']}（{case['respondent_en']}）")
    doc.add_paragraph(case.get("respondent_address_zh", "香港"))
    doc.add_paragraph()
    _body(doc,
        f"現通知你呈請人打算於  {data['hearing_date_zh']}  {data['hearing_time']}，"
        f"在香港灣仔港灣道12號灣仔政府大樓灣仔法院  {data['hearing_floor']}  樓"
        f"第  {data['hearing_room']}  庭內庭，家事法庭  {data.get('hearing_judge', '')}  "
        f"法官申請命令使：")
    doc.add_paragraph()
    _numbered(doc, 1,
        f"呈請人獲得許可，可以於  {data['start_date_zh']}  至  {data['end_date_zh']}  "
        f"帶家庭中的子女  {kids}  離開香港，脫離法庭司法管轄區，"
        f"目的是往  {data['destination_zh']}旅遊；")
    _numbered(doc, 2, "免除答辯人的同意。")
    doc.add_paragraph()
    doc.add_paragraph("日期：      年      月      日")
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph("司法常務官")
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph("答辯人")
    doc.add_paragraph()
    doc.add_paragraph(f"香港身份證編號  {case['respondent_hkid']}")
    return _to_bytes(doc)


def generate_zh_service_affirmation(data):
    """服達誓詞 — Option B (confirms mailing of docs to respondent)."""
    doc = _new_doc()
    case = data.get("case", CASE)
    _zh_court_header(doc, "誓詞", case)
    _body(doc, f"本人  {case['petitioner_zh']}  ，現居於  {case['petitioner_address_zh']}。")
    doc.add_paragraph()
    _body(doc, "謹此謹以至誠確認以下內容：")
    doc.add_paragraph()
    _numbered(doc, 1,
        f"本人於  {data['mail_date_zh']}  在  {data['post_office_zh']}  "
        f"郵寄  傳票/誓詞/承諾書  2份副本到  {data['respondent_address_short_zh']}。")
    _numbered(doc, 2, "而上述文件並無退回。")
    doc.add_paragraph()
    _body(doc, "本人謹此謹以至誠確認，此誓詞內容一概是真。")
    _affirmation_footer_zh(doc, data["sign_year"], data["sign_month"], data["sign_day"])
    return _to_bytes(doc)


# ── English School Trip ───────────────────────────────────────────────────────

def _en_header(doc, title, case):
    for line in [
        "IN THE DISTRICT COURT OF THE",
        "HONG KONG SPECIAL ADMINISTRATIVE REGION",
        f"MATRIMONIAL CAUSES NO. {case['number']} OF {case['year']}",
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(line)
        if "MATRIMONIAL" in line:
            r.bold = True
    doc.add_paragraph("_" * 60)
    doc.add_paragraph()
    doc.add_paragraph("Between")
    pname = doc.add_paragraph()
    pname.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pname.add_run(f"{case['petitioner_en']}（{case['petitioner_zh']}）")
    doc.add_paragraph().alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.paragraphs[-1].add_run("and")
    rname = doc.add_paragraph()
    rname.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rname.add_run(f"{case['respondent_en']}（{case['respondent_zh']}）")
    doc.add_paragraph("_" * 60)
    doc.add_paragraph()
    heading = doc.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = heading.add_run(title.upper())
    r2.bold = True
    r2.underline = True
    doc.add_paragraph()


def generate_en_affirmation(data):
    """Affirmation — English school trip."""
    doc = _new_doc()
    case = data.get("case", CASE)
    child = data["child_en"]
    _en_header(doc, "Affirmation", case)
    _body(doc,
        f"I, {case['petitioner_en']} ({case['petitioner_zh']}), "
        f"of {case['petitioner_address_en']}, "
        f"solemnly and sincerely affirm as follows:")
    doc.add_paragraph()
    _numbered(doc, 1, "I am the Petitioner herein.")
    _numbered(doc, 2,
        f"Custody of the children has been granted to the Petitioner and the "
        f"Respondent pursuant to Order made on {case['custody_order_date_en']}.")
    _numbered(doc, 3,
        f"I hereby apply for taking the child of the family, namely {child}, "
        f"out of Hong Kong for the purpose of attending the school trip to "
        f"{data['destination_en']}, organized by {data['school_name']} "
        f"under the supervision of school teachers, from {data['trip_start_en']} "
        f"to {data['trip_end_en']} (both dates inclusive).")
    _numbered(doc, 4,
        "I annex herewith my undertaking and the Respondent's consent "
        "marked as Exhibit \"A\" and \"B\" respectively.")
    _numbered(doc, 5,
        "I identify the signature of the Respondent in the Consent and I verily "
        "believe that he fully understands the effect and result of such Consent "
        "and that the Consent was given voluntarily.")
    _numbered(doc, 6,
        "And I solemnly and sincerely affirm that the contents of this affirmation are true.")
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph("_" * 30)
    doc.add_paragraph("signature of affirmant")
    doc.add_paragraph()
    doc.add_paragraph(
        f"AFFIRMED at the Court of Justice\n"
        f"in the HKSAR this {data['affirm_date_en']}")
    doc.add_paragraph()
    doc.add_paragraph("Before me,")
    doc.add_paragraph()
    doc.add_paragraph("Commissioner for Oaths\n\t\tJudiciary")
    return _to_bytes(doc)


def generate_en_consent(data):
    """Consent — English school trip (respondent/mother signs)."""
    doc = _new_doc()
    case = data.get("case", CASE)
    child = data["child_en"]
    _en_header(doc, "Consent", case)
    _body(doc,
        f"Whereas the Petitioner is desirous of permitting the child of the family, "
        f"namely {child}, to leave the jurisdiction of Hong Kong for the purpose of "
        f"attending the school trip to {data['destination_en']}, organized by "
        f"{data['school_name']} under the supervision of school teachers, I, the "
        f"Respondent, do hereby give my written consent that the Petitioner may "
        f"permit the said child to leave Hong Kong for such purpose and for such "
        f"period specified below.")
    doc.add_paragraph()
    doc.add_paragraph("The proposed trips are as follows:")
    doc.add_paragraph()
    trip = doc.add_paragraph()
    trip.paragraph_format.left_indent = Cm(0.74)
    r = trip.add_run(
        f"(a)\t{data['destination_en']} – from {data['trip_start_en']} "
        f"to {data['trip_end_en']}.")
    r.bold = True
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    sig = doc.add_paragraph()
    sig.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    sig.add_run("_" * 30)
    np_ = doc.add_paragraph()
    np_.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    np_.add_run(f"{case['respondent_en']} ({case['respondent_zh']})\nRespondent")
    doc.add_paragraph()
    dp = doc.add_paragraph()
    dp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    dp.add_run(f"Dated this {data['affirm_date_en']}")
    return _to_bytes(doc)


def generate_en_undertaking(data):
    """Undertaking — English school trip (petitioner/father signs)."""
    doc = _new_doc()
    case = data.get("case", CASE)
    child = data["child_en"]
    _en_header(doc, "Undertaking", case)
    _body(doc,
        f"I, {case['petitioner_en']} ({case['petitioner_zh']}), am the Petitioner herein.")
    doc.add_paragraph()
    _body(doc,
        f"Being desirous of removing the child of the family, namely {child}, "
        f"from Hong Kong for the purpose of school study trip, I do hereby give my "
        f"undertaking to the Court to return the said child to Hong Kong as and when "
        f"called upon by the Court to do so.")
    doc.add_paragraph()
    doc.add_paragraph("The proposed trip is as follows:")
    doc.add_paragraph()
    trip = doc.add_paragraph()
    trip.paragraph_format.left_indent = Cm(1.27)
    trip.add_run(
        f"{data['destination_en']} – from {data['trip_start_en']} to {data['trip_end_en']}, "
        f"organised by {data['school_name']} under the supervision of school teachers.")
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    sig = doc.add_paragraph()
    sig.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    sig.add_run("_" * 30)
    np_ = doc.add_paragraph()
    np_.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    np_.add_run(f"{case['petitioner_en']} ({case['petitioner_zh']})\nPetitioner")
    doc.add_paragraph()
    dp = doc.add_paragraph()
    dp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    dp.add_run(f"Dated this {data['affirm_date_en']}")
    return _to_bytes(doc)
